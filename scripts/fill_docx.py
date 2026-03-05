"""
Fill Pass/Fail and Actual Response columns in LeadFlow-AI-Testing-Checklist.docx
Based on Newman run: 186/186 assertions passing (100%) on 2026-03-06
"""
import sys
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Mapping: table_row_index -> (pass_fail_symbol, actual_response_text)
# pass_fail: "PASS" | "FAIL" | "N/A"
RESULTS = {
    # 00. Health Check
    2:  ("PASS", '200 OK · {status:"ok", services:{database:"ok", redis:"degraded"}}'),
    3:  ("PASS", '200 OK · {name:"LeadFlow AI CRM API", version:"v1"}'),
    4:  ("FAIL", '404 Not Found · ROUTE_NOT_FOUND — endpoint not implemented'),

    # 01. Auth
    6:  ("PASS", '201 Created · {accessToken, refreshToken, user:{id, email, role}}'),
    7:  ("PASS", '200 OK · {accessToken, refreshToken, user:{id, email, role}}'),
    8:  ("PASS", '200 OK · {id, email, role, organizationId, name}'),
    9:  ("PASS", '200 OK · {accessToken, refreshToken} (token family rotation)'),
    10: ("PASS", '200 OK · {message:"Logged out successfully"} · token blacklisted'),
    11: ("PASS", '401 Unauthorized · INVALID_CREDENTIALS · wrong password rejected'),

    # 02. Projects
    13: ("PASS", '201 Created · {id, name, slug, status, organizationId}'),
    14: ("PASS", '200 OK · {data:[], total, page, limit} (paginated)'),
    15: ("PASS", '200 OK · full project detail with towers'),
    16: ("PASS", '200 OK · updated project fields'),
    17: ("PASS", '200 OK · {total, available, blocked, booked, registered} counts'),
    18: ("PASS", '200 OK · {towers:[{id, name, floors, unitsPerFloor}]}'),

    # 03. Units
    20: ("PASS", '200 OK · paginated list, first 4 unitIds saved as variables'),
    21: ("PASS", '200 OK · unit detail with status, blockExpiresAt, costSheet'),
    22: ("PASS", '200 OK · {basePrice, parkingCharges, otherCharges} in paise'),
    23: ("PASS", '200 OK · {status:"blocked", blockExpiresAt, salesPersonId}'),
    24: ("PASS", '200 OK · 2nd unit blocked by same SP'),
    25: ("PASS", '200 OK · 3rd unit blocked by same SP'),
    26: ("PASS", '422 Unprocessable Entity · BUSINESS_RULE_ERROR: max 3 active blocks per SP'),
    27: ("PASS", '200 OK · unit released, status: available'),

    # 04. Sales Team
    29: ("PASS", '201 Created · {id, name, spCode, role} · salesPersonId saved'),
    30: ("PASS", '200 OK · list with activeBlocks count per SP'),
    31: ("PASS", '200 OK · SP profile with booking stats'),
    32: ("PASS", '200 OK · {achievement%, target, actual} performance'),
    33: ("PASS", '200 OK · designation/target/mobile updated'),

    # 05. Agents
    35: ("PASS", '201 Created · {id, name, agencyName, reraNumber} · agentId saved'),
    36: ("PASS", '200 OK · list with pendingCommission summary'),
    37: ("PASS", '200 OK · {pendingCommission, paidCommission, totalEarned}'),
    38: ("PASS", '200 OK · {rating: 4.5} updated (1–5 scale)'),

    # 06. Customers
    40: ("PASS", '201 Created · customerId saved · PAN masked (ABCDE####F) in response'),
    41: ("PASS", '201 Created · customerId2 saved (for transfer test)'),
    42: ("PASS", '200 OK · {isExisting:true, customer:{id, name, ...}} — no duplicate'),
    43: ("PASS", '400 Bad Request · VALIDATION_ERROR: "must be at least 18 years old"'),
    44: ("PASS", '200 OK · paginated customer list'),
    45: ("PASS", '200 OK · customer detail, PAN masked for non-admin role'),
    46: ("PASS", '200 OK · KYC fields (aadhaar, PAN, address) updated'),

    # 07. Leads
    48: ("PASS", '201 Created · {id, leadCode:"LD-XXXX", status:"new"} · leadId saved'),
    49: ("PASS", '200 OK · paginated list filterable by status/source/SP'),
    50: ("PASS", '200 OK · lead detail with nested site visits + follow-ups'),
    51: ("PASS", '200 OK · {status:"contacted"} — valid transition from new'),
    52: ("PASS", '200 OK · {status:"negotiation"} — valid transition from contacted'),
    53: ("N/A",  'Not in Newman collection — manual test: expects 422 VALIDATION_ERROR'),

    # 08. Site Visits
    55: ("PASS", '201 Created · {id, leadId, scheduledAt, status:"scheduled"} · siteVisitId saved'),
    56: ("PASS", '200 OK · paginated list with status/project/SP filters'),
    57: ("PASS", '200 OK · feedback + followUpDate updated, status: completed'),

    # 09. Follow-Up Tasks
    59: ("PASS", '201 Created · {id, leadId, dueDate, priority} · followUpId saved'),
    60: ("PASS", '200 OK · list filterable by assignedTo/status'),
    61: ("PASS", '200 OK · outcome + nextFollowUpDate set, status: completed'),

    # 10. Bookings
    63: ("PASS", '422 Unprocessable Entity · KYC_INCOMPLETE: missing address, PAN'),
    64: ("PASS", '201 Created · {id, bookingCode, status:"booked"} · bookingId saved'),
    65: ("PASS", '200 OK · paginated list with status/project/customer filters'),
    66: ("PASS", '200 OK · full booking detail with payment schedule breakdown'),
    67: ("PASS", '200 OK · {status:"registered"} · registration cascade triggered'),
    68: ("PASS", '422 Unprocessable Entity · BUSINESS_RULE_ERROR: booking already registered'),

    # 11. Payments
    70: ("PASS", '201 Created · status:"cleared" (NEFT auto-clears) · paymentId + receiptNumber saved'),
    71: ("PASS", '200 OK · {isDuplicate:true, payment:{...}} — same receipt returned'),
    72: ("PASS", '201 Created · status:"under_process" (cheque awaits manual clearance)'),
    73: ("PASS", '200 OK · list of all payments for the booking'),
    74: ("N/A",  'Not in Newman collection — manual test: PATCH status → cleared'),
    75: ("PASS", '200 OK · {status:"bounced"} · complaint auto-created, demand reversed'),
    76: ("N/A",  'Not in Newman collection — manual test: already-cleared → 422'),
    77: ("PASS", '200 OK · {totalPaid, totalPending, pctPaid, breakdown:[]}'),

    # 12. Demand Letters
    79: ("PASS", '201 Created · {id, dueDate, amount} · demandLetterId saved'),
    80: ("PASS", '200 OK · list with status (pending/sent/overdue)'),
    81: ("PASS", '200 OK · full demand letter detail'),
    82: ("PASS", '200 OK · reminder dispatched (email + SMS queued to BullMQ)'),

    # 13. Documents
    84: ("PASS", '200 OK · {documentId, uploadUrl (presigned S3 PUT, 5min expiry)}'),
    85: ("PASS", '200 OK · {status:"uploaded"} after S3 PUT confirmed'),
    86: ("PASS", '200 OK · {downloadUrl (presigned S3 GET, 1hr expiry)}'),
    87: ("PASS", '404 Not Found · pending/unconfirmed document returns 404'),
    88: ("PASS", '201 Created · nocDocumentId saved (direct register, no S3 flow)'),
    89: ("PASS", '201 Created · NOC document registered and confirmed'),
    90: ("PASS", '200 OK · list with category + status filters'),
    91: ("PASS", '200 OK · {status:"verified"} — document marked verified'),

    # 14. Complaints
    93: ("PASS", '201 Created · {id, complaintCode, slaDeadline} · complaintId saved'),
    94: ("PASS", '200 OK · paginated list with status/category/priority filters'),
    95: ("N/A",  'SLA filter not in Newman collection — manual test required'),
    96: ("PASS", '200 OK · full complaint detail with timeline'),
    97: ("PASS", '200 OK · {status:"in_progress"} — assigned + status updated'),
    98: ("PASS", '200 OK · {status:"resolved", resolvedAt} — complaint closed'),

    # 15. Communications
    100: ("PASS", '201 Created · call logged with durationSeconds validated'),
    101: ("PASS", '201 Created · WhatsApp message logged (no duration required)'),
    102: ("N/A",  'Not in Newman collection (only 2 of 3 communication types tested)'),
    103: ("PASS", '200 OK · list of communications for customer'),
    104: ("PASS", '200 OK · activity summary per customer'),

    # 16. Approvals
    106: ("PASS", '200 OK · {pendingCount: N} — UI badge counter'),
    107: ("PASS", '201 Created · approvalId saved · {entityType, status:"pending"}'),
    108: ("PASS", '200 OK · list of all approval requests'),
    109: ("PASS", '200 OK · filtered list with status=pending'),
    110: ("PASS", '200 OK · full approval detail with requestData object'),
    111: ("PASS", '422 Unprocessable · BUSINESS_RULE_ERROR: cannot review own approval (single-user test)'),
    112: ("PASS", '422 Unprocessable · BUSINESS_RULE_ERROR: self-review blocked'),

    # 17. Loans
    114: ("PASS", '201 Created · {id, bankName, sanctionedAmount, status:"applied"} · loanId saved'),
    115: ("PASS", '200 OK · list filterable by status/bank'),
    116: ("PASS", '200 OK · loan detail with disbursement % calculated'),
    117: ("PASS", '200 OK · {status:"sanctioned"} updated'),
    118: ("PASS", '201 Created · disbursement tranche record added to disbursements[]'),

    # 18. Cancellations
    120: ("PASS", '200 OK · {forfeitureAmount, refundableAmount, deductions:[]}'),
    121: ("PASS", '201 Created · cancellationId saved · approval request auto-created'),
    122: ("PASS", '200 OK · list of cancellations with status'),

    # 19. Transfers
    124: ("PASS", '201 Created · transferId saved · approval request auto-created'),
    125: ("PASS", '400 Bad Request · VALIDATION_ERROR: nocDocumentId required for transfer'),
    126: ("PASS", '200 OK · list of transfers with status'),

    # 20. Possessions
    128: ("PASS", '200 OK · possessionId saved, list filterable by status/project'),
    129: ("PASS", '200 OK · full possession detail'),
    130: ("PASS", '200 OK · scheduledDate + offerLetterSentAt updated'),
    131: ("PASS", '201 Created · snagId saved · {description, severity, location}'),
    132: ("PASS", '200 OK · snag list filterable by severity/status'),
    133: ("PASS", '200 OK · {status:"resolved"} — snag resolved'),
    134: ("PASS", '200 OK · possession completed · cascade: sale_completed + commissions triggered'),

    # 21. Analytics
    136: ("PASS", '200 OK · {cached:false} — fresh data from DB'),
    137: ("PASS", '200 OK · {cached:true/false} — cached field is boolean'),
    138: ("PASS", '200 OK · revenue + bookings breakdown by SP/project/month'),
    139: ("PASS", '200 OK · payments + overdue + breakdown by payment mode'),

    # 22. Audit Log
    141: ("PASS", '200 OK · paginated audit log (PAN/Aadhaar never in metadata)'),
    142: ("PASS", '200 OK · full entity trail for booking with all actions'),
    143: ("PASS", '200 OK · all actions performed by user'),

    # 23. Security & Edge Cases
    145: ("PASS", '401 Unauthorized · AUTH_REQUIRED — missing token rejected'),
    146: ("PASS", '401 Unauthorized · INVALID_TOKEN — malformed JWT rejected'),
    147: ("PASS", '404 Not Found · ROUTE_NOT_FOUND — unknown route handled'),
    148: ("PASS", '200 OK · public endpoint — no auth token required'),
    149: ("N/A",  'Rate-limit headers not explicitly asserted in Newman collection'),
}

SYMBOL = {
    "PASS": "✅",
    "FAIL": "❌",
    "N/A":  "➖",
}


def set_cell_text(cell, text, bold=False, color=None):
    """Clear cell and set text, preserving paragraph formatting."""
    para = cell.paragraphs[0]
    # Clear existing runs
    for run in para.runs:
        run.text = ""
    # Add new run or reuse first
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run()
    run.text = text
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)


def main():
    path = 'reports/LeadFlow-AI-Testing-Checklist.docx'
    doc = Document(path)
    table = doc.tables[2]

    pass_count = 0
    fail_count = 0
    na_count = 0

    for row_idx, row in enumerate(table.rows):
        if row_idx not in RESULTS:
            continue

        result, actual = RESULTS[row_idx]
        symbol = SYMBOL[result]

        # Col 4: Pass/Fail
        pf_text = f"{symbol} {result}"
        if result == "PASS":
            set_cell_text(row.cells[4], pf_text, bold=True, color=(0, 128, 0))
            pass_count += 1
        elif result == "FAIL":
            set_cell_text(row.cells[4], pf_text, bold=True, color=(192, 0, 0))
            fail_count += 1
        else:
            set_cell_text(row.cells[4], pf_text, bold=False, color=(128, 128, 128))
            na_count += 1

        # Col 5: Actual Response
        set_cell_text(row.cells[5], actual)

    # Update Table 1 (coverage summary) — add totals row if needed
    # (leave as-is, just save)

    out_path = 'reports/LeadFlow-AI-Testing-Checklist-filled.docx'
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"  PASS: {pass_count}")
    print(f"  FAIL: {fail_count}")
    print(f"  N/A:  {na_count}")
    print(f"  Total rows filled: {pass_count + fail_count + na_count}")


if __name__ == "__main__":
    main()
